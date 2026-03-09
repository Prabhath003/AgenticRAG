"""File utilities for type detection and conversion."""

from typing import Literal, Optional, Dict, Tuple, Any, List
import magic
import subprocess
import tempfile
import os
import io
import mimetypes
from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Table,
    Spacer,
)
from reportlab.lib.units import inch

# Programming language detection based on file extensions
EXTENSION_TO_LANGUAGE: Dict[str, str] = {
    ".py": "python",
    ".js": "javascript",
    ".ts": "typescript",
    ".tsx": "typescript",
    ".jsx": "javascript",
    ".java": "java",
    ".cpp": "cpp",
    ".c": "c",
    ".cs": "csharp",
    ".rb": "ruby",
    ".go": "go",
    ".rs": "rust",
    ".php": "php",
    ".swift": "swift",
    ".kt": "kotlin",
    ".scala": "scala",
    ".sh": "bash",
    ".bash": "bash",
    ".zsh": "zsh",
    ".fish": "fish",
    ".ps1": "powershell",
    ".sql": "sql",
    ".html": "html",
    ".css": "css",
    ".scss": "scss",
    ".sass": "sass",
    ".less": "less",
    ".xml": "xml",
    ".json": "json",
    ".yaml": "yaml",
    ".yml": "yaml",
    ".toml": "toml",
    ".ini": "ini",
    ".cfg": "ini",
    ".conf": "conf",
    ".md": "markdown",
    ".tex": "latex",
    ".r": "r",
    ".R": "r",
    ".jl": "julia",
    ".pl": "perl",
    ".lua": "lua",
    ".vim": "vim",
    ".gradle": "gradle",
    ".maven": "maven",
    ".dockerfile": "dockerfile",
    ".Dockerfile": "dockerfile",
    ".makefile": "makefile",
    ".Makefile": "makefile",
    ".txt": "plaintext",
    ".text": "plaintext",
    ".log": "plaintext",
}

# Fallback MIME type mappings for edge cases
CONTENT_TYPE_MAP: Dict[str, str] = {
    ".md": "text/markdown",
    ".sh": "text/x-shellscript",
    ".ts": "text/typescript",
}

# Constants for optimal model context balance
MAX_FILE_SIZE = 20 * 1024  # 20 KB
MAX_DIR_ITEMS_PER_LEVEL = 30

# MIME type to file type mapping
MIME_TYPE_MAPPING: Dict[str, Literal["image", "pdf", "docx", "xlsx", "ppt", "text", "markdown"]] = {
    "image/": "image",
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.openxmlformats-officedocument.presentationml.presentation": "ppt",
    "text/": "text",
}

# File extensions for specific types
EXTENSION_MAPPING: Dict[Tuple[str, str], Literal["markdown"]] = {
    (".md", ".markdown"): "markdown",
}


def get_content_type(filename: str) -> str:
    """
    Detect MIME type from filename using mimetypes module.

    Uses Python's built-in mimetypes database which covers thousands of file types,
    with fallback to custom mappings for special cases.

    Args:
        filename: The name of the file (with extension)

    Returns:
        MIME type string (e.g., "image/png", "application/pdf")
    """
    # Try standard mimetypes module first (covers ~1000+ types)
    mime_type, _ = mimetypes.guess_type(filename)
    if mime_type:
        return mime_type

    # Fall back to custom mappings for edge cases
    file_ext = Path(filename).suffix.lower()
    if file_ext in CONTENT_TYPE_MAP:
        return CONTENT_TYPE_MAP[file_ext]

    # Default to binary/octet-stream for unknown types
    return "application/octet-stream"


def detect_file_type(
    filename: str,
    content_type: str,
    file_content: Optional[bytes] = None,
    extension_to_language: Optional[Dict[str, str]] = None,
) -> Literal["code", "unknown", "docx", "xlsx", "image", "pdf", "markdown", "text", "ppt"]:
    """
    Detect the file type from filename, content_type, and file content using python-magic.

    Uses mapping tables for maintainability and extensibility.

    Args:
        filename: The name of the file
        content_type: The MIME type from headers
        file_content: Optional binary file content for accurate detection
        extension_to_language: Optional dict mapping file extensions to language names

    Returns: 'image', 'pdf', 'markdown', 'code', 'text', 'docx', 'xlsx', 'ppt', or 'unknown'
    """
    try:
        # If file content is available, use magic for accurate detection
        if file_content:
            mime_type = magic.from_buffer(file_content, mime=True)
        else:
            mime_type = content_type
    except (ImportError, Exception):
        # Fallback if python-magic is not available
        mime_type = content_type

    lower_name = filename.lower()

    # Check MIME type mapping (exact and prefix matches)
    for mime_pattern, file_type in MIME_TYPE_MAPPING.items():
        if mime_pattern.endswith("/"):  # Prefix match (e.g., "image/")
            if mime_type.startswith(mime_pattern):
                return file_type
        else:  # Exact match
            if mime_type == mime_pattern:
                return file_type

    # Check extension mapping
    for extensions, file_type in EXTENSION_MAPPING.items():
        if lower_name.endswith(extensions):
            return file_type

    # Code files - check against extension_to_language if provided
    if extension_to_language:
        for ext in extension_to_language.keys():
            if lower_name.endswith(ext):
                return "code"

    return "unknown"


def text_to_markdown(text_content: str, filename: str, is_code: bool = False) -> str:
    """
    Convert plain text or code to markdown format.

    Args:
        text_content: The text content to convert
        filename: The filename (used to detect language)
        is_code: Whether to treat as code block

    Returns:
        Markdown formatted text
    """
    if is_code:
        # Get language from filename extension
        lower_name = filename.lower()
        language = "text"
        for ext, lang in EXTENSION_TO_LANGUAGE.items():
            if lower_name.endswith(ext):
                language = lang
                break
        return f"```{language}\n{text_content}\n```"
    else:
        # Return plain text as-is for markdown
        return text_content


def office_document_to_pdf(content: bytes, filename: str, doc_type: str) -> bytes:
    """
    Convert office documents (DOCX, XLSX, PPTX) to PDF.

    Args:
        content: File content bytes
        filename: The filename
        doc_type: 'docx', 'xlsx', or 'ppt'

    Returns:
        PDF content as bytes, or original content if conversion fails
    """
    try:
        # Try using subprocess with libreoffice if available

        with tempfile.NamedTemporaryFile(suffix=Path(filename).suffix, delete=False) as tmp_input:
            tmp_input.write(content)
            tmp_input_path = tmp_input.name

        with tempfile.TemporaryDirectory() as tmp_dir:
            try:
                # Use libreoffice to convert to PDF
                subprocess.run(
                    [
                        "libreoffice",
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        tmp_dir,
                        tmp_input_path,
                    ],
                    timeout=30,
                    capture_output=True,
                    check=True,
                )

                # Get the output PDF file
                pdf_filename = Path(tmp_input_path).stem + ".pdf"
                pdf_path = os.path.join(tmp_dir, pdf_filename)

                if os.path.exists(pdf_path):
                    with open(pdf_path, "rb") as pdf_file:
                        pdf_content = pdf_file.read()
                    return pdf_content

            except (
                subprocess.CalledProcessError,
                FileNotFoundError,
                TimeoutError,
            ):
                pass  # Continue to fallback

            finally:
                if os.path.exists(tmp_input_path):
                    os.remove(tmp_input_path)

        # Fallback: Try python-pptx/python-docx with reportlab
        if doc_type == "docx":
            try:
                doc = Document(io.BytesIO(content))
                pdf_buffer = io.BytesIO()
                pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
                elements: List[Any] = []
                styles = getSampleStyleSheet()

                for para in doc.paragraphs:
                    if para.text.strip():
                        elements.append(Paragraph(para.text, styles["Normal"]))

                for table in doc.tables:
                    table_data: List[List[Any]] = []
                    for row in table.rows:
                        table_data.append([cell.text for cell in row.cells])
                    if table_data:
                        elements.append(Table(table_data))
                        elements.append(Spacer(1, 0.2 * inch))

                if elements:
                    pdf_doc.build(elements)
                    pdf_buffer.seek(0)
                    return pdf_buffer.read()

            except (ImportError, Exception):
                pass  # Continue to return original content

    except Exception:
        pass  # Continue to return original content

    # Return original content if conversion fails
    return content
